import sys
print("PYTHON EXE:", sys.executable)
print("PYTHON VERSION:", sys.version)

# ======================================================================
# IMPORTS
# ======================================================================
import os
import re
import json
import time
import shutil
import sqlite3
from datetime import datetime
from functools import lru_cache

import fitz  # noqa: F401 – used elsewhere
import pandas as pd
from docx import Document
from docx.shared import Pt  # noqa: F401
from flask import (Flask, render_template, request, redirect,
                   url_for, send_file, flash)

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "hard_to_guess_string")

# ======================================================================
# BASIC WORKING STRUCTURE OF THE WEBSITE
# Excel upload →
#   template_configs/template_id.json
#   template_databases/template_id.db
#   template_documents/template_id/
#       ├─ template_id.docx   (Word template – placeholder वाला)
#       └─ template_id.pdf    (Uploaded / base PDF)
#   generated_docs/template_id/
#       └─ template_id_x.docx (records से बना हुआ)
# ======================================================================

# ----------------------------------------------------------------------
# DIRECTORIES
# ----------------------------------------------------------------------
CONFIG_DIR        = "template_configs"
DB_DIR            = "template_databases"
DOC_TEMPLATE_DIR  = "template_documents"
GENERATED_DIR     = "generated_docs"
UPLOAD_BASE_DIR   = "uploaded_pdf"

for _d in (CONFIG_DIR, DB_DIR, DOC_TEMPLATE_DIR, GENERATED_DIR, UPLOAD_BASE_DIR):
    os.makedirs(_d, exist_ok=True)

# Tesseract path – override via env var for cross-platform support
pytesseract_cmd = os.getenv(
    "TESSERACT_CMD",
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

# ----------------------------------------------------------------------
# SQL RESERVED KEYWORDS SET
# ----------------------------------------------------------------------
SQL_RESERVED_KEYWORDS = {
    "abort", "action", "add", "after", "all", "alter", "analyze", "and",
    "as", "asc", "attach", "autoincrement", "before", "begin", "between",
    "by", "cascade", "case", "cast", "check", "collate", "column", "commit",
    "conflict", "constraint", "create", "cross", "current_date",
    "current_time", "current_timestamp", "database", "default", "deferrable",
    "deferred", "delete", "desc", "detach", "distinct", "drop", "each",
    "else", "end", "escape", "except", "exclusive", "exists", "explain",
    "fail", "for", "foreign", "from", "full", "glob", "group", "having",
    "if", "ignore", "immediate", "in", "index", "indexed", "initially",
    "inner", "insert", "instead", "intersect", "into", "is", "isnull",
    "join", "key", "left", "like", "limit", "match", "natural", "no",
    "not", "notnull", "null", "of", "offset", "on", "or", "order",
    "outer", "plan", "pragma", "primary", "query", "raise", "recursive",
    "references", "regexp", "reindex", "release", "rename", "replace",
    "restrict", "right", "rollback", "row", "savepoint", "select", "set",
    "table", "temp", "temporary", "then", "to", "transaction", "trigger",
    "union", "unique", "update", "using", "vacuum", "values", "view",
    "virtual", "when", "where", "with", "without",
}

# ----------------------------------------------------------------------
# SAFE FILENAME / SQL IDENTIFIER HELPERS
# ----------------------------------------------------------------------
def sanitize_column_name(name):
    """Replace non-word chars with underscore; avoid SQL reserved words."""
    base = re.sub(r'[^\w]', '_', str(name)).lower()
    if base in SQL_RESERVED_KEYWORDS:
        base += "_col"
    return base


def make_unique_columns(columns):
    """Return a list of sanitized, unique column names."""
    seen = {}
    result = []
    for col in columns:
        s = sanitize_column_name(str(col))
        if s in seen:
            seen[s] += 1
            s = f"{s}_{seen[s]}"
        else:
            seen[s] = 0
        result.append(s)
    return result

# ----------------------------------------------------------------------
# LOAD TEMPLATE CONFIG
# ----------------------------------------------------------------------
def load_template_config(template_id):
    path = os.path.join(CONFIG_DIR, f"{template_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template config not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

# ----------------------------------------------------------------------
# INIT DB IF NEEDED
# ----------------------------------------------------------------------
def init_db_if_needed(config):
    db_path = os.path.join(DB_DIR, config["database"])
    columns_def = ", ".join(f"'{f['name']}' TEXT" for f in config["fields"])
    with sqlite3.connect(db_path, check_same_thread=False) as conn:
        conn.execute(f"""
            CREATE TABLE IF NOT EXISTS records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                {columns_def}
            )
        """)
        conn.commit()
    return db_path

# ----------------------------------------------------------------------
# LOAD MANUAL TEMPLATES (cached)
# ----------------------------------------------------------------------
@lru_cache(maxsize=1)
def load_manual_templates():
    """
    Home page aur sidebar ke liye templates load karta hai.
    Sirf auto-generated pending templates skip hote hain
    (jinmein 'source_query' hai — ye dashboard se accessible hain).
    """
    templates = []
    for file in os.listdir(CONFIG_DIR):
        if file.endswith(".json"):
            with open(os.path.join(CONFIG_DIR, file), "r", encoding="utf-8") as f:
                config = json.load(f)
            if "source_query" in config:
                continue
            templates.append(config)
    return templates

# ----------------------------------------------------------------------
# CATEGORIES WITH UPLOADED FILES
# ----------------------------------------------------------------------
@lru_cache(maxsize=1)
def get_categories_with_files():
    """Return list of categories that have at least one uploaded file."""
    result = []
    if os.path.exists(UPLOAD_BASE_DIR):
        for category in os.listdir(UPLOAD_BASE_DIR):
            category_path = os.path.join(UPLOAD_BASE_DIR, category)
            if os.path.isdir(category_path):
                files = [
                    f for f in os.listdir(category_path)
                    if os.path.isfile(os.path.join(category_path, f))
                ]
                if files:
                    result.append({"name": category, "count": len(files)})
    return sorted(result, key=lambda x: x["name"])

# ----------------------------------------------------------------------
# CONTEXT PROCESSOR
# ----------------------------------------------------------------------
@app.context_processor
def inject_templates():
    try:
        templates = load_manual_templates()
    except Exception:
        templates = []
    return dict(
        templates=templates,
        categories_with_files=get_categories_with_files(),
    )

# ======================================================================
# PLACEHOLDER REPLACEMENT
# ======================================================================
def replace_placeholders(doc, record, placeholders):

    def get_value(field, case=None, extract=None):
        value = str(record.get(field, "") or "")
        if extract:
            m = re.search(extract, value)
            value = m.group(1) if m else value
        if case == "upper":
            value = value.upper()
        return value

    def replace_in_runs(paragraph):
        text = "".join(run.text for run in paragraph.runs)
        updated_text = text

        for placeholder, mapping in placeholders.items():
            if isinstance(mapping, dict):
                field = mapping.get("field")
                case  = mapping.get("case")
                multi = mapping.get("multi_paragraph", False)

                # AUTO DATE PLACEHOLDER
                if field is None and mapping.get("type") == "auto_date":
                    fmt = (mapping.get("format", "DD.MM.YYYY")
                           .replace("DD", "%d").replace("MM", "%m").replace("YYYY", "%Y"))
                    replacement = datetime.now().strftime(fmt)
                    updated_text = updated_text.replace(placeholder, replacement)
                    continue

                if field is None:
                    continue

                extract = mapping.get("extract")
            else:
                field   = mapping
                case    = None
                multi   = False
                extract = None

            replacement = get_value(field, case, extract)

            # APPLY REPLACE LOGIC
            if isinstance(mapping, dict) and "replace" in mapping:
                for old, new in mapping["replace"].items():
                    replacement = replacement.replace(old, new)

            # MULTI PARAGRAPH HANDLING (only once, before text replace)
            if multi and placeholder in text:
                lines = replacement.split("\n")
                paragraph.text = lines[0]
                for line in lines[1:]:
                    new_p = paragraph.insert_paragraph_before(line)
                    new_p.style = paragraph.style
                return

            updated_text = updated_text.replace(placeholder, replacement)

        # APPLY UPDATED TEXT
        if updated_text != text:
            for r in paragraph.runs:
                r.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = updated_text
            else:
                paragraph.add_run(updated_text)

    def walk(doc):
        for p in doc.paragraphs:
            replace_in_runs(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_runs(p)
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_in_runs(p)
            for p in section.footer.paragraphs:
                replace_in_runs(p)
            if section.different_first_page_header_footer:
                for p in section.first_page_header.paragraphs:
                    replace_in_runs(p)
                for p in section.first_page_footer.paragraphs:
                    replace_in_runs(p)

    walk(doc)
    return doc

# ======================================================================
# HOME / TEMPLATES
# ======================================================================
@app.route("/")
def home():
    templates = load_manual_templates()
    return render_template("home.html", templates=templates)


@app.route("/form/<template_id>")
def form(template_id):
    try:
        config = load_template_config(template_id)
        init_db_if_needed(config)
        return render_template("index.html", config=config, template_id=template_id)
    except Exception as e:
        flash(f"Error loading template: {e}")
        return redirect(url_for("home"))

# ----------------------------------------------------------------------
# ADD TEMPLATE FROM EXCEL
# ----------------------------------------------------------------------
@app.route("/add_template", methods=["GET", "POST"])
def add_template():
    if request.method == "POST":
        file         = request.files.get("file")
        display_name = request.form.get("display_name")

        if not file or not display_name:
            flash("Please provide both Excel file and template name")
            return redirect(url_for("add_template"))

        try:
            df = pd.read_excel(file)
            if df.empty:
                flash("Excel file is empty")
                return redirect(url_for("add_template"))

            original_cols  = list(df.columns)
            sanitized_cols = make_unique_columns(original_cols)
            df.columns     = sanitized_cols

            base_name   = sanitize_column_name(display_name.lower().replace(" ", "_"))
            template_id = base_name

            template_doc_dir       = os.path.join(DOC_TEMPLATE_DIR, template_id)
            generated_template_dir = os.path.join(GENERATED_DIR, template_id)
            os.makedirs(template_doc_dir, exist_ok=True)
            os.makedirs(generated_template_dir, exist_ok=True)

            db_filename       = f"{template_id}.db"
            json_filename     = f"{template_id}.json"
            doc_template_name = f"{template_id}.docx"

            fields = []
            template_placeholders = {}
            for i, col in enumerate(original_cols):
                sc = sanitized_cols[i]
                fields.append({"name": sc, "label": str(col), "type": "text",
                               "placeholder": f"Enter {col}"})
                template_placeholders[f"<{sc.upper()}>"] = sc

            template_json = {
                "id": template_id,
                "display_name": display_name,
                "database": db_filename,
                "document_template": f"{template_id}/{doc_template_name}",
                "base_pdf": f"{template_id}/{template_id}.pdf",
                "macro_file": "sample.bas",
                "export_options": {
                    "export_excel": True,
                    "export_word": True,
                    "excel_filename": f"{template_id}.xlsx",
                    "word_template": f"{template_id}/{doc_template_name}",
                },
                "fields": fields,
                "placeholders": template_placeholders,
            }

            with open(os.path.join(CONFIG_DIR, json_filename), "w", encoding="utf-8") as f:
                json.dump(template_json, f, indent=4)

            db_path = init_db_if_needed(template_json)
            with sqlite3.connect(db_path, check_same_thread=False) as conn:
                for _, row in df.iterrows():
                    vals = [str(row[col]) if pd.notna(row[col]) else ""
                            for col in sanitized_cols]
                    conn.execute(
                        f"INSERT INTO records ({','.join(f'\"{c}\"' for c in sanitized_cols)}) "
                        f"VALUES ({','.join(['?'] * len(sanitized_cols))})",
                        vals,
                    )
                conn.commit()

            # Auto-generate Word template
            doc_path = os.path.join(template_doc_dir, doc_template_name)
            doc = Document()
            doc.add_heading(display_name, level=1)
            doc.add_paragraph("<< AUTO-GENERATED TEMPLATE >>")
            for ph in template_placeholders:
                p   = doc.add_paragraph()
                run = p.add_run(ph)
                run.bold = True
            doc.save(doc_path)

            load_manual_templates.cache_clear()
            flash(f"Template '{display_name}' added successfully!")
            return redirect(url_for("home"))

        except Exception as e:
            flash(f"Error adding template: {e}")
            return redirect(url_for("add_template"))

    return render_template("add_template.html")

# ======================================================================
# DELETE TEMPLATE
# ======================================================================
@app.route("/delete_template/<template_id>", methods=["POST"])
def delete_template(template_id):
    try:
        config_path = os.path.join(CONFIG_DIR, f"{template_id}.json")
        if os.path.exists(config_path):
            os.remove(config_path)

        db_path = os.path.join(DB_DIR, f"{template_id}.db")
        try:
            sqlite3.connect(db_path, check_same_thread=False).close()
        except Exception:
            pass

        if os.path.exists(db_path):
            deadline = time.time() + 5
            while True:
                try:
                    os.remove(db_path)
                    break
                except PermissionError:
                    if time.time() > deadline:
                        raise
                    time.sleep(0.1)

        template_folder = os.path.join(DOC_TEMPLATE_DIR, template_id)
        if os.path.exists(template_folder):
            shutil.rmtree(template_folder)

        load_manual_templates.cache_clear()
        flash(f"Template '{template_id}' deleted successfully!")
    except Exception as e:
        flash(f"Error deleting template: {e}")

    return redirect(url_for("home"))

# ======================================================================
# RECORD MANAGEMENT
# ======================================================================
@app.route("/submit/<template_id>", methods=["POST"])
def submit(template_id):
    try:
        config  = load_template_config(template_id)
        db_path = os.path.join(DB_DIR, config["database"])
        fields  = config["fields"]
        data    = [request.form.get(f["name"], "") for f in fields]

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.execute(
                f"INSERT INTO records ({','.join(f['name'] for f in fields)}) "
                f"VALUES ({','.join(['?'] * len(fields))})",
                data,
            )
            conn.commit()
        flash("Record saved successfully!")
        return redirect(url_for("records", template_id=template_id))
    except Exception as e:
        flash(f"Error saving: {e}")
        return redirect(url_for("form", template_id=template_id))



# ======================================================================
# PAGINATION HELPER — shared by records, all_records, pending_records
# ======================================================================
def _build_query_parts(request_args, allowed_cols=None):
    """
    Read GET params and return (where_extras, params, sort_col,
    sort_dir, col_filters, page, search).
    allowed_cols: whitelist for sort/column-filter col names (SQL injection guard).
    """
    import re as _re
    page     = request_args.get("page", 1, type=int)
    search   = request_args.get("q", "").strip()
    sort_col = request_args.get("sort", "").strip()
    sort_dir = request_args.get("dir", "asc").strip().lower()
    if sort_dir not in ("asc", "desc"):
        sort_dir = "asc"

    # Whitelist sort column
    if allowed_cols and sort_col not in allowed_cols:
        sort_col = ""

    # Per-column filters: col_<name>=value
    col_filters = {}
    for k, v in request_args.items():
        if k.startswith("col_") and v.strip():
            col = k[4:]
            if allowed_cols is None or col in allowed_cols:
                col_filters[col] = v.strip()

    extras  = []
    params  = []

    # Global search
    if search and allowed_cols:
        like_parts = [f"CAST(\"{c}\" AS TEXT) LIKE ?" for c in allowed_cols]
        extras.append("(" + " OR ".join(like_parts) + ")")
        params    += [f"%{search}%"] * len(allowed_cols)

    # Per-column filter
    for col, val in col_filters.items():
        extras.append(f"CAST(\"{col}\" AS TEXT) LIKE ?")
        params.append(f"%{val}%")

    return extras, params, sort_col, sort_dir, col_filters, page, search

@app.route("/records/<template_id>")
def records(template_id):
    try:
        import math
        config   = load_template_config(template_id)
        PER_PAGE = 20

        db_path = (os.path.join(DB_DIR, config["database"])
                   if "source_query" in config
                   else init_db_if_needed(config))

        base_sql = (f"SELECT * FROM ({config['source_query']}) AS base"
                    if "source_query" in config
                    else "SELECT * FROM records")

        # Get all column names for whitelist
        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.row_factory = sqlite3.Row
            sample = conn.execute(f"{base_sql} LIMIT 1").fetchone()
            all_cols = list(sample.keys()) if sample else []

        extras, params, sort_col, sort_dir, col_filters, page, search = \
            _build_query_parts(request.args, allowed_cols=all_cols)

        where_clause = ("WHERE " + " AND ".join(extras)) if extras else ""
        order_clause = (f'ORDER BY "{sort_col}" {sort_dir.upper()}' if sort_col
                        else "ORDER BY id DESC")

        count_sql = f"SELECT COUNT(*) FROM ({base_sql}) AS t {where_clause}"
        data_sql  = (f"SELECT * FROM ({base_sql}) AS t {where_clause}"
                     f" {order_clause} LIMIT {PER_PAGE} OFFSET {(page-1)*PER_PAGE}")

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.execute("PRAGMA journal_mode=WAL")
            conn.row_factory = sqlite3.Row
            total = conn.execute(count_sql, params).fetchone()[0]
            rows  = conn.execute(data_sql,  params).fetchall()

        return render_template("records.html",
                               data=rows, config=config, template_id=template_id,
                               page=page,
                               total_pages=max(1, math.ceil(total / PER_PAGE)),
                               total=total, per_page=PER_PAGE,
                               search=search, sort_col=sort_col,
                               sort_dir=sort_dir, col_filters=col_filters,
                               all_cols=all_cols)
    except Exception as e:
        import traceback; traceback.print_exc()
        flash(f"Error loading records: {e}")
        return redirect(url_for("home"))


@app.route("/edit/<template_id>/<int:record_id>", methods=["GET", "POST"])
def edit_record(template_id, record_id):
    try:
        config  = load_template_config(template_id)
        db_path = os.path.join(DB_DIR, config["database"])
        fields  = config["fields"]

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.row_factory = sqlite3.Row
            record = conn.execute(
                "SELECT * FROM records WHERE id=?", (record_id,)
            ).fetchone()

        if not record:
            flash("Record not found")
            return redirect(url_for("records", template_id=template_id))

        record_dict = {key: (record[key] or "") for key in record.keys()}

        if request.method == "POST":
            updated    = [request.form.get(f["name"], "") for f in fields]
            set_clause = ", ".join(f"{f['name']}=?" for f in fields)
            with sqlite3.connect(db_path, check_same_thread=False) as conn:
                conn.execute(
                    f"UPDATE records SET {set_clause} WHERE id=?",
                    updated + [record_id],
                )
                conn.commit()
            flash("Updated!")
            return redirect(url_for("records", template_id=template_id))

        return render_template(
            "edit_record.html", record=record_dict,
            config=config, template_id=template_id,
        )
    except Exception as e:
        flash(f"Edit error: {e}")
        return redirect(url_for("records", template_id=template_id))


@app.route("/delete/<template_id>/<int:record_id>")
def delete_record(template_id, record_id):
    try:
        config  = load_template_config(template_id)
        db_path = os.path.join(DB_DIR, config["database"])
        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.execute("DELETE FROM records WHERE id=?", (record_id,))
            conn.commit()
        flash("Deleted!")
    except Exception as e:
        flash(f"Delete error: {e}")
    return redirect(url_for("records", template_id=template_id))

# ======================================================================
# EXCEL IMPORT / EXPORT
# ======================================================================
@app.route("/import_excel/<template_id>", methods=["POST"])
def import_excel(template_id):
    file = request.files.get("file")
    if not file:
        flash("No file selected")
        return redirect(url_for("records", template_id=template_id))
    try:
        action      = request.form.get("action")
        config      = load_template_config(template_id)
        field_names = [f["name"] for f in config["fields"]]
        db_path     = os.path.join(DB_DIR, config["database"])
        df          = pd.read_excel(file)
        missing     = [c for c in field_names if c not in df.columns]
        if missing:
            flash(f"Missing columns: {', '.join(missing)}")
            return redirect(url_for("records", template_id=template_id))

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            if action == "overwrite":
                conn.execute("DELETE FROM records")
            for _, row in df.iterrows():
                vals = [str(row[c]) if pd.notna(row[c]) else "" for c in field_names]
                conn.execute(
                    f"INSERT INTO records ({','.join(field_names)}) "
                    f"VALUES ({','.join(['?'] * len(field_names))})",
                    vals,
                )
            conn.commit()
        flash("Imported successfully!")
    except Exception as e:
        flash(f"Import error: {e}")
    return redirect(url_for("records", template_id=template_id))


@app.route("/export_word/<template_id>")
def export_word(template_id):
    try:
        config   = load_template_config(template_id)
        db_path  = os.path.join(DB_DIR, config["database"])
        out_file = os.path.join(GENERATED_DIR,
                                f"{sanitize_column_name(template_id)}_records.docx")

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.row_factory = sqlite3.Row
            all_records = conn.execute("SELECT * FROM records").fetchall()

        doc = Document()
        doc.add_heading(f"{config['display_name']} Records", level=1)
        for rec in all_records:
            for field in config["fields"]:
                doc.add_paragraph(f"{field['label']}: {rec[field['name']]}")
            doc.add_paragraph("-" * 23)
        doc.save(out_file)
        return send_file(out_file, as_attachment=True)
    except Exception as e:
        flash(f"Word export error: {e}")
        return redirect(url_for("records", template_id=template_id))


@app.route("/export_excel/<template_id>")
def export_excel(template_id):
    try:
        config  = load_template_config(template_id)
        db_path = os.path.join(DB_DIR, config["database"])
        out     = os.path.join(GENERATED_DIR,
                               f"{sanitize_column_name(template_id)}_records.xlsx")
        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            df = pd.read_sql_query("SELECT * FROM records", conn)
            df.to_excel(out, index=False)
        return send_file(out, as_attachment=True)
    except Exception as e:
        flash(f"Export error: {e}")
        return redirect(url_for("records", template_id=template_id))

# ======================================================================
# WORD DOCUMENT GENERATION
# ======================================================================
@app.route("/generate_doc/<template_id>/<int:record_id>")
def generate_doc(template_id, record_id):
    try:
        config  = load_template_config(template_id)
        db_path = os.path.join(DB_DIR, config["database"])

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.row_factory = sqlite3.Row
            record = conn.execute(
                "SELECT * FROM records WHERE id=?", (record_id,)
            ).fetchone()

        if not record:
            flash("Record not found")
            return redirect(url_for("records", template_id=template_id))

        template_file = config.get("document_template")
        doc_id        = request.args.get("doc")

        if doc_id and config.get("download_options"):
            for opt in config["download_options"]:
                if opt["id"] == doc_id:
                    template_file = opt["template"]
                    break

        template_path = os.path.join(DOC_TEMPLATE_DIR, template_file)
        if not os.path.exists(template_path):
            flash(f"Template file not found: {template_file}")
            return redirect(url_for("records", template_id=template_id))

        output_dir  = os.path.join(GENERATED_DIR, template_id)
        os.makedirs(output_dir, exist_ok=True)
        safe_doc_id = doc_id or "default"
        output_path = os.path.join(output_dir,
                                   f"{template_id}_{record_id}_{safe_doc_id}.docx")

        doc         = Document(template_path)
        record_data = dict(record)

        rem_no = request.args.get("rem_no")
        if rem_no:
            record_data["rem_no"] = rem_no

        for key in config.get("runtime_inputs", []):
            value = request.args.get(key)
            if not value:
                flash(f"Please provide {key.upper()}")
                return redirect(url_for("records", template_id=template_id))
            record_data[key] = value

        try:
            ph_config = config.get("placeholders", {})
            if isinstance(ph_config, str):
                ph_config = json.loads(ph_config)
            replace_placeholders(doc, record_data, ph_config)
        except Exception as e:
            flash(f"Generation error: {e}")
            return redirect(url_for("records", template_id=template_id))

        doc.save(output_path)
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        print(f"Unhandled error: {e}")
        flash(f"Generation error: {e}")
        return redirect(url_for("records", template_id=template_id))

# ======================================================================
# DASHBOARD – CATEGORY QUERIES, FY HELPERS, STATS
# ======================================================================

# ──────────────────────────────────────────────────────────────────────
# STEP 1 — MASTER CATEGORY DEFINITION TABLE
#
# Each entry defines ONE category with:
#   sql   : WHERE clause fragment used directly in SQLite queries
#   pandas: lambda(df) → filtered DataFrame — derived from same logic,
#           used only in _compute_sidebar_stats (in-memory pass)
#
# "Misc" is intentionally absent here; it is derived at runtime as
# NOT (union of all sql conditions below).
#
# To add/change a category: edit ONLY this table. Everything else
# (SQL queries, pandas filters, CATEGORY_LIST, normalized lookups,
# MISC condition, sidebar stats) auto-updates from here.
# ──────────────────────────────────────────────────────────────────────
_CATEGORY_DEFS = [
    # ──────────────────────────────────────────────────────────────────
    # Each entry has:
    #   name   : display name (used in CATEGORY_LIST, URLs, dashboard)
    #   sql    : SQLite WHERE fragment  (used in SQL queries)
    #   pandas : lambda(df)→df  (used in in-memory stats pass)
    #   sync   : list of dicts — each dict describes ONE target DB to
    #            populate during sync_category_databases().
    #            Keys per dict:
    #              db      : filename inside DB_DIR
    #              filter  : (optional) override lambda — use when the
    #                        DB needs a tighter/different slice than the
    #                        category's own pandas fn (e.g. RTI morly-only,
    #                        or Ca3-M owning the combined A+B+M slice)
    #              lang    : (optional) "ENG" or "HIN" — adds a
    #                        lang==value sub-filter before writing
    #            No "cols" needed — category DB column names now match
    #            master DB column names exactly, so the filtered slice
    #            is written as-is (master columns that aren't in the
    #            category DB schema are simply ignored by SQLite).
    #            If sync is absent/None → category is dashboard-only
    #            (no individual DB to populate).
    # ──────────────────────────────────────────────────────────────────
    {
        "name": "Announcement",
        "sql":  "(LOWER(puc) LIKE '%announcement%' OR LOWER(file) LIKE '%/pas/%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("announcement", na=False, regex=False) |
            df["file"].str.lower().str.contains("/pas/", na=False, regex=False)
        ],
        "sync": [{"db": "announcement.db"}],
    },
    {
        "name": "Booth",
        "sql":  "(LOWER(puc) LIKE '%booth%' OR LOWER(puc) LIKE '%helpdesk%' OR LOWER(puc) LIKE '%help-desk%' "
                "OR LOWER(file) LIKE '%booth%' OR LOWER(file) LIKE '%helpdesk%' OR LOWER(file) LIKE '%help-desk%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("booth|helpdesk|help-desk", na=False) |
            df["file"].str.lower().str.contains("booth|helpdesk|help-desk", na=False)
        ],
        "sync": [{"db": "booth.db"}],
    },
    {
        "name": "BOV",
        "sql":  "(LOWER(puc) LIKE '%battery%' OR LOWER(puc) LIKE '%bov%' OR LOWER(puc) LIKE '%boc%' "
                "OR LOWER(file) LIKE '%bov%' OR LOWER(file) LIKE '%boc%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("battery|bov|boc", na=False) |
            df["file"].str.lower().str.contains("bov|boc", na=False)
        ],
        "sync": None,   # dashboard-only; no individual category DB
    },
    {
        "name": "Display",
        "sql":  "(LOWER(puc) LIKE '%board%' OR LOWER(puc) LIKE '%display%' "
                "OR LOWER(file) LIKE '%/cib/%' OR LOWER(file) LIKE '%db%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("board|display", na=False) |
            df["file"].str.lower().str.contains("/cib/|db", na=False)
        ],
        "sync": None,
    },
    {
        "name": "Coordination",
        "sql":  "LOWER(file) LIKE '%coord%'",
        "pandas": lambda df: df[
            df["file"].str.lower().str.contains("coord", na=False, regex=False)
        ],
        "sync": [{"db": "coordination_pm_wing.db"}],
    },
    {
        "name": "Sahayaks",
        "sql":  "(LOWER(puc) LIKE '%sahayak%' OR LOWER(puc) LIKE '%coolie%' "
                "OR LOWER(file) LIKE '%sahayak%' OR LOWER(file) LIKE '%coolie%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("sahayak|coolie", na=False) |
            df["file"].str.lower().str.contains("sahayak|coolie", na=False)
        ],
        "sync": None,
    },
    {
        "name": "Union",
        "sql":  "(LOWER(puc) LIKE '%pnm%' OR LOWER(puc) LIKE '%jcm%' OR LOWER(puc) LIKE '%union%' "
                "OR LOWER(file) LIKE '%pnm%' OR LOWER(file) LIKE '%jcm%' OR LOWER(file) LIKE '%union%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("pnm|jcm|union", na=False) |
            df["file"].str.lower().str.contains("pnm|jcm|union", na=False)
        ],
        "sync": None,
    },
    {
        "name": "Parl",
        "sql":  "(LOWER(case_col) LIKE 'ls/%' OR LOWER(case_col) LIKE 'rs/%' OR LOWER(case_col) LIKE '%r377/%' "
                "OR LOWER(case_col) LIKE '%zh/%' OR LOWER(case_col) LIKE '%assurance/%' OR LOWER(case_col) LIKE '%sm/%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("ls/|rs/|r377/|zh/|assurance/|sm/", na=False)
        ],
        "sync": [{"db": "parliament.db"}],
    },
    {
        "name": "Comm",
        "sql":  "LOWER(from_col) LIKE '%d/f(bc)%'",
        "pandas": lambda df: df[
            df["from_col"].str.lower().str.contains("d/f(bc)", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "RTI",
        "sql":  "(LOWER(puc) LIKE '%rti app%' OR LOWER(file) LIKE '%/rti/%' "
                "OR LOWER(case_col) LIKE '%morly%' OR LOWER(case_col) LIKE '%rti%')",
        "pandas": lambda df: df[
            df["puc"].str.lower().str.contains("rti app", na=False, regex=False) |
            df["file"].str.lower().str.contains("/rti/", na=False, regex=False) |
            df["case_col"].str.lower().str.contains("morly|rti", na=False)
        ],
        # RTI DB uses a tighter filter (morly only) — override pandas with a
        # custom sub-filter lambda via the optional "filter" key
        "sync": [
            {
                "db": "rti.db",
                "filter": lambda df: df[
                    df["case_col"].str.lower().str.contains("morly", na=False, regex=False)
                ],
            }
        ],
    },
    {
        "name": "CPGRAM",
        "sql":  "(LOWER(case_col) LIKE '%pmopg%' OR LOWER(case_col) LIKE '%prsec%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("pmopg|prsec", na=False)
        ],
        "sync": None,
    },
    {
        "name": "Court",
        "sql":  "(LOWER(case_col) LIKE '%court%' OR LOWER(case_col) LIKE 'ccpwd/%' OR LOWER(case_col) LIKE 'wp%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("court|ccpwd|wp", na=False)
        ],
        "sync": None,
    },
    {
        "name": "NHRC",
        "sql":  "LOWER(case_col) LIKE '%nhrc%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("nhrc", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "PMO",
        "sql":  "LOWER(case_col) LIKE '%pmo%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("pmo", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "Ca3-A",
        "sql":  "(LOWER(case_col) = 'a' OR LOWER(case_col) LIKE 'a/%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.match(r"^a(/.*)?$", na=False)
        ],
        # Ca3-A belongs to the combined Representations group (A+B+M)
        # The combined filter + lang split is handled in Ca3-M (last of the three)
        # to avoid writing 3× the same rows.  Individual entries: sync=None.
        "sync": None,
    },
    {
        "name": "Ca3-B",
        "sql":  "(LOWER(case_col) = 'b' OR LOWER(case_col) LIKE 'b/%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.match(r"^b(/.*)?$", na=False)
        ],
        "sync": None,
    },
    {
        "name": "Ca3-C",
        "sql":  "(LOWER(case_col) = 'c' OR LOWER(case_col) LIKE 'c/%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.match(r"^c(/.*)?$", na=False)
        ],
        # Ca3-C → References (ENG + HIN)
        "sync": [
            {"db": "ref_eng.db", "lang": "ENG"},
            {"db": "ref_hin.db", "lang": "HIN"},
        ],
    },
    {
        "name": "Ca3-M",
        "sql":  "(LOWER(case_col) = 'm' OR LOWER(case_col) LIKE 'm/%')",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.match(r"^m(/.*)?$", na=False)
        ],
        # Ca3-M is the last of A/B/M group → owns the combined Representations DBs.
        # A custom "filter" lambda unions all three (A, B, M) before the lang split.
        "sync": [
            {
                "db": "ref_eng_rep.db",
                "lang": "ENG",
                "filter": lambda df: df[
                    df["case_col"].str.lower().str.match(r"^[abm](/.*)?$", na=False)
                ],
            },
            {
                "db": "ref_hin_rep.db",
                "lang": "HIN",
                "filter": lambda df: df[
                    df["case_col"].str.lower().str.match(r"^[abm](/.*)?$", na=False)
                ],
            },
        ],
    },
    {
        "name": "PQ-LS",
        "sql":  "LOWER(case_col) LIKE 'ls/pq/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("ls/pq/", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "PQ-RS",
        "sql":  "LOWER(case_col) LIKE 'rs/pq/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("rs/pq/", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "Rule-377",
        "sql":  "LOWER(case_col) LIKE '%r377/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("r377/", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "Zero-Hour",
        "sql":  "LOWER(case_col) LIKE '%zh/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("zh/", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "Assurance",
        "sql":  "LOWER(case_col) LIKE '%assurance/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("assurance/", na=False, regex=False)
        ],
        "sync": None,
    },
    {
        "name": "Spl-Mention",
        "sql":  "LOWER(case_col) LIKE '%sm/%'",
        "pandas": lambda df: df[
            df["case_col"].str.lower().str.contains("sm/", na=False, regex=False)
        ],
        "sync": None,
    },
]

# ──────────────────────────────────────────────────────────────────────
# STEP 2 — AUTO-DERIVED GLOBALS  (never edit these manually)
#
# Everything below is computed once at startup from _CATEGORY_DEFS.
# ──────────────────────────────────────────────────────────────────────

# Ordered list of display names (Misc appended last)
CATEGORY_LIST = [d["name"] for d in _CATEGORY_DEFS] + ["Misc"]

# {display_name → sql_condition}  — replaces old CATEGORY_QUERIES dict
CATEGORY_QUERIES = {d["name"]: d["sql"] for d in _CATEGORY_DEFS}

# {normalized_name → sql_condition}  — used by _resolve_category_condition
# pre-built once so we never re-iterate CATEGORY_QUERIES per request
_NORM_QUERIES = {
    d["name"].lower().replace(" ", "_").replace("-", "_").replace("/", "_"): d["sql"]
    for d in _CATEGORY_DEFS
}

# {normalized_name → pandas_filter_fn}  — used by _filter_category
_NORM_PANDAS = {
    d["name"].lower().replace(" ", "_").replace("-", "_").replace("/", "_"): d["pandas"]
    for d in _CATEGORY_DEFS
}

# Pre-built MISC SQL condition  =  NOT (union of all 24 category conditions)
_ALL_CATEGORIES_UNION_SQL = " OR ".join(f"({d['sql']})" for d in _CATEGORY_DEFS)
_MISC_SQL_CONDITION        = f"NOT ({_ALL_CATEGORIES_UNION_SQL})"

# SQLite date-filter expression — one place to change the date format
_DATE_FILTER = "DATE(substr(dt,7,4)||'-'||substr(dt,4,2)||'-'||substr(dt,1,2))"


# ──────────────────────────────────────────────────────────────────────
# STEP 3 — UTILITY FUNCTIONS that use the globals above
# ──────────────────────────────────────────────────────────────────────

def normalize_category(category):
    """'Ca3-A' → 'ca3_a',  'PQ-LS' → 'pq_ls',  etc."""
    return category.lower().replace(" ", "_").replace("-", "_").replace("/", "_")


def all_categories_union_sql(category_queries=None):
    """Return pre-built union SQL (arg kept for backward-compat but ignored)."""
    return _ALL_CATEGORIES_UNION_SQL


def _filter_category(df, category):
    """
    Filter a DataFrame to rows matching `category`.
    Looks up the pre-compiled pandas lambda from _NORM_PANDAS.
    Returns full df for 'Misc' (caller handles misc logic separately).
    """
    key = normalize_category(category)
    fn  = _NORM_PANDAS.get(key)
    if fn is None:
        return df          # Misc — caller must handle
    return fn(df)


def _resolve_category_condition(category_normalized):
    """
    Return (sql_condition, is_valid) for a normalized category name.
    Uses pre-built _NORM_QUERIES and _MISC_SQL_CONDITION — no iteration.
    """
    if category_normalized == "misc":
        return _MISC_SQL_CONDITION, True
    cond = _NORM_QUERIES.get(category_normalized)
    return cond, cond is not None


# ──────────────────────────────────────────────
# FY BOUNDS
# ──────────────────────────────────────────────
def get_fy_bounds(today=None):
    if today is None:
        today = datetime.now()

    y = today.year
    if today.month >= 4:
        cur_start  = datetime(y, 4, 1)
        cur_end    = datetime(y + 1, 3, 31)
        prev_start = datetime(y - 1, 4, 1)
        prev_end   = datetime(y, 3, 31)
        cur_label  = f"{y}-{str(y + 1)[-2:]}"
        prev_label = f"{y - 1}-{str(y)[-2:]}"
    else:
        cur_start  = datetime(y - 1, 4, 1)
        cur_end    = datetime(y, 3, 31)
        prev_start = datetime(y - 2, 4, 1)
        prev_end   = datetime(y - 1, 3, 31)
        cur_label  = f"{y - 1}-{str(y)[-2:]}"
        prev_label = f"{y - 2}-{str(y - 1)[-2:]}"

    return dict(
        cur_start=cur_start, cur_end=cur_end,
        prev_start=prev_start, prev_end=prev_end,
        cur_label=cur_label, prev_label=prev_label,
    )


# ──────────────────────────────────────────────
# GLOBAL SIDEBAR STATS HELPER
# Single DB query → pandas → all 25 category counts in one pass.
# ──────────────────────────────────────────────
def _compute_sidebar_stats(db_path, fy_bounds):
    """
    Returns list of dicts:
        { category, current_fy_total, current_fy_pending,
          previous_fy_total, previous_fy_pending }
    One SELECT, everything else in-memory.
    """
    with sqlite3.connect(db_path, check_same_thread=False) as conn:
        df = pd.read_sql_query(
            "SELECT status, puc, file, case_col, from_col, dt FROM records",
            conn,
        )

    # Parse dates once
    df["date"] = pd.to_datetime(
        df["dt"].str[6:10] + "-" + df["dt"].str[3:5] + "-" + df["dt"].str[0:2],
        errors="coerce",
    )

    cur_rows  = df[(df["date"] >= fy_bounds["cur_start"])  & (df["date"] <= fy_bounds["cur_end"])]
    prev_rows = df[(df["date"] >= fy_bounds["prev_start"]) & (df["date"] <= fy_bounds["prev_end"])]

    def _get_misc(rows):
        matched = set()
        for d in _CATEGORY_DEFS:                         # use defs directly
            matched |= set(d["pandas"](rows).index)
        return rows[~rows.index.isin(matched)]

    stats = []
    for category in CATEGORY_LIST:
        if category == "Misc":
            c_rows = _get_misc(cur_rows)
            p_rows = _get_misc(prev_rows)
        else:
            c_rows = _filter_category(cur_rows,  category)
            p_rows = _filter_category(prev_rows, category)

        stats.append({
            "category":            category,
            "current_fy_total":    len(c_rows),
            "current_fy_pending":  len(c_rows[c_rows["status"].str.lower().str.contains("pending", na=False)]),
            "previous_fy_total":   len(p_rows),
            "previous_fy_pending": len(p_rows[p_rows["status"].str.lower().str.contains("pending", na=False)]),
        })
    return stats


# ──────────────────────────────────────────────
# HELPER: load master DB path from config
# ──────────────────────────────────────────────
def _get_master_db_path():
    with open(os.path.join(CONFIG_DIR, "master_status.json")) as f:
        return os.path.join(DB_DIR, json.load(f)["database"])


# ──────────────────────────────────────────────
# DASHBOARD
# ──────────────────────────────────────────────
@app.route("/dashboard")
def dashboard():
    try:
        master_config_path = os.path.join(CONFIG_DIR, "master_status.json")
        if not os.path.exists(master_config_path):
            flash("master_status.json not found!")
            return redirect(url_for("home"))

        with open(master_config_path, "r", encoding="utf-8") as f:
            master_config = json.load(f)

        master_db_path = os.path.join(DB_DIR, master_config["database"])
        if not os.path.exists(master_db_path):
            flash("master_status.db not found!")
            return redirect(url_for("home"))

        fy    = get_fy_bounds()
        stats = _compute_sidebar_stats(master_db_path, fy)

        return render_template(
            "dashboard.html",
            stats=stats,
            current_fy=fy["cur_label"],
            previous_fy=fy["prev_label"],
            normalize_category=normalize_category,
            categories=CATEGORY_LIST,
        )
    except Exception as e:
        flash(f"Error loading dashboard: {e}")
        return redirect(url_for("home"))


# ──────────────────────────────────────────────
# ALL RECORDS VIEW
# ──────────────────────────────────────────────
@app.route("/all_records/<category>/<fy>")
def all_records(category, fy):
    """Show ALL records (pending + completed) for a category and FY."""
    import math
    try:
        cat_norm = normalize_category(category)
        db_path  = _get_master_db_path()

        condition, valid = _resolve_category_condition(cat_norm)
        if not valid:
            flash("Invalid category")
            return redirect(url_for("dashboard"))

        fy_bounds = get_fy_bounds()
        if fy == "current":
            start_date, end_date, fy_label = (
                fy_bounds["cur_start"], fy_bounds["cur_end"], fy_bounds["cur_label"])
        else:
            start_date, end_date, fy_label = (
                fy_bounds["prev_start"], fy_bounds["prev_end"], fy_bounds["prev_label"])

        PER_PAGE   = 20
        MASTER_COLS = ["puc", "file", "case_col", "from_col", "dt", "status"]

        base_fixed = (f"{condition}"
                      f" AND {_DATE_FILTER} BETWEEN DATE('{start_date}') AND DATE('{end_date}')")

        extras, qparams, sort_col, sort_dir, col_filters, page, search = \
            _build_query_parts(request.args, allowed_cols=MASTER_COLS)

        full_where   = "WHERE " + base_fixed
        if extras:
            full_where += " AND " + " AND ".join(extras)

        order_clause = (f'ORDER BY "{sort_col}" {sort_dir.upper()}' if sort_col
                        else "ORDER BY dt DESC")

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.execute("PRAGMA journal_mode=WAL")
            total = conn.execute(
                f"SELECT COUNT(*) FROM records {full_where}", qparams
            ).fetchone()[0]
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                f"SELECT * FROM records {full_where}"
                f" {order_clause} LIMIT {PER_PAGE} OFFSET {(page-1)*PER_PAGE}",
                qparams,
            ).fetchall()

        stats = _compute_sidebar_stats(db_path, fy_bounds)

        return render_template(
            "dashboard.html",
            records=[dict(r) for r in rows],
            detail_mode=True,
            category=category.replace("_", " ").title(),
            record_type="All",
            fy_label=fy_label,
            stats=stats,
            page=page,
            total_pages=max(1, math.ceil(total / PER_PAGE)),
            total=total, search=search,
            sort_col=sort_col, sort_dir=sort_dir,
            col_filters=col_filters, all_cols=MASTER_COLS,
            category_raw=category, fy=fy,
        )
    except Exception as e:
        flash(f"Error loading records: {e}")
        return redirect(url_for("dashboard"))


# ──────────────────────────────────────────────
# PENDING RECORDS VIEW
# ──────────────────────────────────────────────
@app.route("/pending_records/<category>/<fy>")
def pending_records(category, fy):
    """Show PENDING records for a category and FY."""
    import math
    try:
        cat_norm = normalize_category(category)
        db_path  = _get_master_db_path()

        condition, valid = _resolve_category_condition(cat_norm)
        if not valid:
            flash("Invalid category")
            return redirect(url_for("dashboard"))

        fy_bounds = get_fy_bounds()
        if fy == "current":
            start_date, end_date, fy_label = (
                fy_bounds["cur_start"], fy_bounds["cur_end"], fy_bounds["cur_label"])
        else:
            start_date, end_date, fy_label = (
                fy_bounds["prev_start"], fy_bounds["prev_end"], fy_bounds["prev_label"])

        PER_PAGE    = 20
        MASTER_COLS = ["puc", "file", "case_col", "from_col", "dt", "status"]

        base_fixed = (f"LOWER(status) LIKE '%pending%'"
                      f" AND {condition}"
                      f" AND {_DATE_FILTER} BETWEEN DATE('{start_date}') AND DATE('{end_date}')")

        extras, qparams, sort_col, sort_dir, col_filters, page, search = \
            _build_query_parts(request.args, allowed_cols=MASTER_COLS)

        full_where   = "WHERE " + base_fixed
        if extras:
            full_where += " AND " + " AND ".join(extras)

        order_clause = (f'ORDER BY "{sort_col}" {sort_dir.upper()}' if sort_col
                        else "ORDER BY dt DESC")

        with sqlite3.connect(db_path, check_same_thread=False) as conn:
            conn.execute("PRAGMA journal_mode=WAL")
            total = conn.execute(
                f"SELECT COUNT(*) FROM records {full_where}", qparams
            ).fetchone()[0]
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                f"SELECT * FROM records {full_where}"
                f" {order_clause} LIMIT {PER_PAGE} OFFSET {(page-1)*PER_PAGE}",
                qparams,
            ).fetchall()

        stats = _compute_sidebar_stats(db_path, fy_bounds)

        return render_template(
            "dashboard.html",
            records=[dict(r) for r in rows],
            detail_mode=True,
            category=category.replace("_", " ").title(),
            record_type="Pending",
            fy_label=fy_label,
            stats=stats,
            page=page,
            total_pages=max(1, math.ceil(total / PER_PAGE)),
            total=total, search=search,
            sort_col=sort_col, sort_dir=sort_dir,
            col_filters=col_filters, all_cols=MASTER_COLS,
            category_raw=category, fy=fy,
        )
    except Exception as e:
        flash(f"Error loading records: {e}")
        return redirect(url_for("dashboard"))


# ======================================================================
# SYNC MASTER STATUS
# ======================================================================
@app.route("/sync_master_status")
def sync_master_status():
    try:
        excel_path         = r"E:\official\03 trfc gnrl\18th ls\status\status.xls"
        master_config_path = os.path.join(CONFIG_DIR, "master_status.json")

        if not os.path.exists(master_config_path):
            flash("❌ master_status.json not found!")
            return redirect(request.referrer or url_for("home"))

        with open(master_config_path, "r", encoding="utf-8") as f:
            master_config = json.load(f)

        master_db_path = os.path.join(DB_DIR, master_config["database"])

        if not os.path.exists(excel_path):
            flash(f"❌ Excel file not found at: {excel_path}")
            return redirect(request.referrer or url_for("home"))

        df = pd.read_excel(excel_path)

        # Clean column names
        df.columns = (
            df.columns.str.strip().str.lower()
              .str.replace(" ", "_").str.replace(r"[^a-z0-9_]", "", regex=True)
        )

        # Convert date columns to dd.mm.yyyy
        date_cols = [c for c in df.columns if "dt" in c or "date" in c]
        for col in date_cols:
            df[col] = pd.to_datetime(df[col], errors="coerce").dt.strftime("%d.%m.%Y")

        with sqlite3.connect(master_db_path, check_same_thread=False) as conn:
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA synchronous=NORMAL")
            # Stable IDs ke liye: poori table replace karo lekin
            # rowid Excel ke natural order se mile (1,2,3...)
            conn.execute("DELETE FROM records")
            conn.commit()
            # index=True, index_label="id" se Excel ka 1-based row number
            # hi id ban jaata hai — har sync pe same Excel = same IDs
            df_indexed = df.copy()
            df_indexed.index = range(1, len(df_indexed) + 1)
            df_indexed.to_sql("records", conn, if_exists="append",
                              index=True, index_label="id")
            rows_imported = len(df)

        sync_category_databases(df)

        flash(f"✅ Successfully synced {rows_imported} records from Excel to all databases!")
        return redirect(request.referrer or url_for("home"))

    except Exception as e:
        import traceback; traceback.print_exc()
        flash(f"❌ Error syncing data: {str(e)}")
        return redirect(request.referrer or url_for("home"))


def sync_category_databases(master_df):
    """
    Sync pending records from master_df into individual category databases.

    Fully data-driven from _CATEGORY_DEFS["sync"] metadata.

    For each category def that has a non-None "sync" list:
      1. Start with pending_df  (status == pending)
      2. Apply "filter" lambda if present, else use category's own pandas fn
      3. If "lang" key present, further filter by lang == value
      4. RECREATE target DB and INSERT filtered slice
    """
    try:
        # ── Step 0: pending slice ────────────────────────────────────────
        if "status" in master_df.columns:
            pending_df = master_df[
                master_df["status"].str.lower().str.strip() == "pending"
            ].copy()
        else:
            pending_df = master_df.copy()

        # ── Helper: write slice to DB ────────────────────────────────────
        def _write(db_name, slice_df):
            db_path = os.path.join(DB_DIR, db_name)
            if os.path.exists(db_path):
                os.remove(db_path)
            with sqlite3.connect(db_path) as conn:
                conn.execute("PRAGMA journal_mode=WAL")
                conn.execute("PRAGMA synchronous=NORMAL")
                if not slice_df.empty:
                    # Master DB ka original 'id' preserve karo —
                    # agar 'id' column hai toh use as index write karo
                    if "id" in slice_df.columns:
                        write_df = slice_df.set_index("id")
                        write_df.to_sql("records", conn, if_exists="replace",
                                        index=True, index_label="id")
                    else:
                        # fallback: 1-based index
                        slice_df = slice_df.copy()
                        slice_df.index = range(1, len(slice_df) + 1)
                        slice_df.to_sql("records", conn, if_exists="replace",
                                        index=True, index_label="id")
                else:
                    conn.execute("""
                        CREATE TABLE IF NOT EXISTS records (
                            id INTEGER PRIMARY KEY AUTOINCREMENT
                        )
                    """)
                conn.commit()

        # ── Main loop ────────────────────────────────────────────────────
        for cat_def in _CATEGORY_DEFS:
            sync_list = cat_def.get("sync")
            if not sync_list:
                continue

            for sync_spec in sync_list:

                filter_fn = sync_spec.get("filter") or cat_def["pandas"]

                base_slice = filter_fn(pending_df)

                lang_val = sync_spec.get("lang")
                if lang_val and "lang" in base_slice.columns:
                    base_slice = base_slice[
                        base_slice["lang"].str.upper() == lang_val
                    ]

                _write(sync_spec["db"], base_slice)

        print("✅ Category databases synced successfully!")

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"❌ Error syncing category databases: {e}")

# ======================================================================
# PDF / FILE UPLOAD (category-wise)
# ======================================================================

@app.route("/upload_category/<category>", methods=["GET", "POST"])
def upload_category(category):
    try:
        category = category.strip()
        if not category:
            flash("Category missing!")
            return redirect(url_for("dashboard"))

        uploaded_files_dir = os.path.join(UPLOAD_BASE_DIR, category)
        os.makedirs(uploaded_files_dir, exist_ok=True)

        if request.method == "POST" and "files[]" in request.files:
            uploaded_files = request.files.getlist("files[]")
            uploaded_names = []
            for file in uploaded_files:
                if file and file.filename:
                    file.save(os.path.join(uploaded_files_dir, file.filename))
                    uploaded_names.append(file.filename)
            if uploaded_names:
                flash(f"Uploaded successfully: {', '.join(uploaded_names)}")
            get_categories_with_files.cache_clear()
            return redirect(url_for("upload_category", category=category))

        files   = sorted(f for f in os.listdir(uploaded_files_dir)
                         if os.path.isfile(os.path.join(uploaded_files_dir, f)))
        remarks = {f: "" for f in files}

        return render_template(
            "upload_files.html",
            template_id=category,
            template_name=category,
            files=files,
            remarks=remarks,
            query="",
            search_results=[],
            ocr_status=None,
        )
    except Exception as e:
        flash(f"Error: {e}")
        return redirect(url_for("dashboard"))


@app.route("/view_file/<category>/<filename>")
def view_uploaded_file(category, filename):
    try:
        file_path = os.path.join(UPLOAD_BASE_DIR, category, filename)
        if not os.path.exists(file_path):
            flash("File not found!")
            return redirect(url_for("upload_category", category=category))

        ext = filename.rsplit(".", 1)[-1].lower()
        mime_types = {
            "pdf":  "application/pdf",
            "jpg":  "image/jpeg", "jpeg": "image/jpeg",
            "png":  "image/png",  "gif":  "image/gif",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "txt":  "text/plain",
        }
        return send_file(file_path, mimetype=mime_types.get(ext, "application/octet-stream"))
    except Exception as e:
        flash(f"Error viewing file: {e}")
        return redirect(url_for("upload_category", category=category))


@app.route("/download_file/<category>/<filename>")
def download_uploaded_file(category, filename):
    try:
        file_path = os.path.join(UPLOAD_BASE_DIR, category, filename)
        if not os.path.exists(file_path):
            flash("File not found!")
            return redirect(url_for("upload_category", category=category))
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f"Error downloading file: {e}")
        return redirect(url_for("upload_category", category=category))


@app.route("/delete_file/<category>/<filename>", methods=["POST"])
def delete_uploaded_file(category, filename):
    try:
        file_path = os.path.join(UPLOAD_BASE_DIR, category, filename)
        if os.path.exists(file_path):
            os.remove(file_path)
            flash(f"File '{filename}' deleted successfully!")
            get_categories_with_files.cache_clear()
        else:
            flash("File not found!")
    except Exception as e:
        flash(f"Error deleting file: {e}")
    return redirect(url_for("upload_category", category=category))


# ======================================================================
# OCR
# ======================================================================

def init_ocr_db():
    with sqlite3.connect("ocr.db") as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS ocr_data (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                category    TEXT,
                filename    TEXT,
                page_number INTEGER,
                content     TEXT
            )
        """)
        conn.commit()

init_ocr_db()


def _do_ocr(file_path, ext):
    """Return list of (page_number, text) tuples."""
    import pytesseract
    from PIL import Image
    pytesseract.pytesseract.tesseract_cmd = pytesseract_cmd

    if ext == "pdf":
        from pdf2image import convert_from_path
        pages = convert_from_path(file_path)
        return [(i + 1, pytesseract.image_to_string(p)) for i, p in enumerate(pages)]
    if ext in {"jpg", "jpeg", "png", "gif"}:
        return [(1, pytesseract.image_to_string(Image.open(file_path)))]
    return []


@app.route("/run_ocr/<category>/<filename>", methods=["POST"])
def run_ocr_file(category, filename):
    try:
        file_path = os.path.join(UPLOAD_BASE_DIR, category, filename)
        if not os.path.exists(file_path):
            flash("File not found for OCR!")
            return redirect(url_for("upload_category", category=category))

        ext    = filename.rsplit(".", 1)[-1].lower()
        pages  = _do_ocr(file_path, ext)
        if not pages:
            flash("Unsupported file type for OCR.")
            return redirect(url_for("upload_category", category=category))

        with sqlite3.connect("ocr.db") as conn:
            conn.execute(
                "DELETE FROM ocr_data WHERE category=? AND filename=?",
                (category, filename),
            )
            conn.executemany(
                "INSERT INTO ocr_data (category, filename, page_number, content) VALUES (?,?,?,?)",
                [(category, filename, pg, txt) for pg, txt in pages],
            )
            conn.commit()

        flash(f"OCR completed for {filename}!")
    except Exception as e:
        flash(f"OCR error: {e}")
    return redirect(url_for("upload_category", category=category))


@app.route("/run_ocr_all/<category>", methods=["POST"])
def run_ocr_all(category):
    uploaded_files_dir = os.path.join(UPLOAD_BASE_DIR, category)
    if not os.path.exists(uploaded_files_dir):
        flash("Category not found!")
        return redirect(url_for("dashboard"))

    files = [
        f for f in os.listdir(uploaded_files_dir)
        if os.path.isfile(os.path.join(uploaded_files_dir, f))
    ]

    with sqlite3.connect("ocr.db") as conn:
        for filename in files:
            ext   = filename.rsplit(".", 1)[-1].lower()
            pages = _do_ocr(os.path.join(uploaded_files_dir, filename), ext)
            conn.execute(
                "DELETE FROM ocr_data WHERE category=? AND filename=?",
                (category, filename),
            )
            if pages:
                conn.executemany(
                    "INSERT INTO ocr_data (category, filename, page_number, content) VALUES (?,?,?,?)",
                    [(category, filename, pg, txt) for pg, txt in pages],
                )
        conn.commit()

    flash("OCR completed for all files!")
    return redirect(url_for("upload_category", category=category))


@app.route("/search/<category>", methods=["POST"])
def search_files(category):
    query   = request.form.get("query", "").strip()
    results = []

    if query:
        with sqlite3.connect("ocr.db") as conn:
            rows = conn.execute(
                "SELECT filename, page_number FROM ocr_data "
                "WHERE category=? AND content LIKE ?",
                (category, f"%{query}%"),
            ).fetchall()
        results = [{"filename": r[0], "page_number": r[1]} for r in rows]

    return render_template(
        "upload_files.html",
        template_id=category,
        template_name=category,
        files=sorted(os.listdir(os.path.join(UPLOAD_BASE_DIR, category))),
        query=query,
        search_results=results,
        ocr_status=None,
    )


# ======================================================================
# RUN
# ======================================================================
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5001)


# from waitress import serve
# if __name__ == "__main__":
#     serve(app, host="0.0.0.0", port=5001)
